"""Build ИНСТРУКЦИЯ.docx for end users."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "ИНСТРУКЦИЯ.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def main() -> None:
    doc = Document()
    title = doc.add_heading("Media Monitor — инструкция для пользователя", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_para(doc, "Версия программы: v1.0.2", bold=True)
    add_para(
        doc,
        "Программа ищет заданные слова и фразы на указанных сайтах СМИ "
        "и сохраняет результат в Excel-отчёт.",
    )

    add_heading(doc, "1. Что лежит в папке", 1)
    add_bullets(
        doc,
        [
            "media-monitor.exe — запуск программы",
            "sites.txt — список сайтов для проверки",
            "words.txt — список слов и фраз для поиска",
            "settings.txt — настройки (даты, лимиты)",
            "exclude.txt — ссылки, которые не должны попадать в отчёт",
            "*.example.txt — образцы файлов (можно смотреть, править не обязательно)",
            "reports\\ — папка с Excel-отчётами",
            "media-monitor_log.txt — лог работы (то же, что в консоли)",
            "ИНСТРУКЦИЯ.docx — эта инструкция",
        ],
    )
    add_para(doc, "Папку можно копировать на флешку целиком — привязки к полному пути нет.")

    add_heading(doc, "2. Быстрый старт", 1)
    add_numbered(
        doc,
        [
            "Откройте sites.txt — по одной ссылке на сайт в строке (например gubdaily.ru).",
            "Откройте words.txt — по одному слову или фразе в строке (например сбертройка).",
            "При необходимости настройте settings.txt (см. раздел 4).",
            "Запустите media-monitor.exe.",
            "Если спросят логин/пароль — можно нажать Enter (пропуск) или ввести данные.",
            "Дождитесь окончания. В конце будет «Нажмите Enter…».",
            "Откройте свежий файл в папке reports\\.",
        ],
    )

    add_heading(doc, "3. Как работает поиск", 1)
    add_bullets(
        doc,
        [
            "Регистр не важен: СберТройка и сбертройка — одно и то же.",
            "Строка в words.txt ищется целиком как фраза (можно несколько слов).",
            "Кавычки в тексте статьи не мешают: «тройка» тоже найдётся.",
            "Ищется текст статьи, а не меню/подвал сайта.",
            "Программа раскрывает списки новостей и использует поиск сайта, чтобы найти статьи.",
        ],
    )

    add_heading(doc, "4. Настройки (settings.txt)", 1)
    add_bullets(
        doc,
        [
            "article_date_not_later_than — дата ДО (включительно). Пусто = сегодня.",
            "article_date_not_older_than — дата НЕ СТАРШЕ. Пример 2026-01-01 — только с этой даты и новее. Пусто = без нижней границы.",
            "auth_timeout_seconds — сколько секунд ждать логин/пароль при старте.",
            "max_scan_urls — максимум страниц за запуск. 0 = без лимита.",
            "max_expand_links — сколько ссылок брать с одной страницы списка/поиска.",
            "ssl_verify — 1 проверять сертификаты, 0 не проверять (быстрее на госсайтах с битым SSL).",
        ],
    )
    add_para(
        doc,
        "Если дата статьи найдена и она вне диапазона — статья не попадёт в отчёт. "
        "Если дата не найдена — статья всё равно попадёт (ячейка даты будет пустой).",
    )

    add_heading(doc, "5. Исключения (exclude.txt)", 1)
    add_para(
        doc,
        "Если статья находится, но в отчёт её не нужно — добавьте её полный URL "
        "(по одной ссылке в строке).",
    )

    add_heading(doc, "6. Отчёт Excel", 1)
    add_para(doc, "Файлы: reports\\media-monitor-ГГГГ-ММ-ДД_ЧЧ-ММ-СС.xlsx")
    add_bullets(
        doc,
        [
            "A — слово / фраза",
            "B — ссылка на страницу (кликабельная)",
            "C — дата выхода статьи",
            "D — название статьи",
            "E — дата/время скана",
            "F1 — версия программы (например v1.0.2)",
            "G1 — длительность генерации отчёта (например 12м 35с)",
        ],
    )

    add_heading(doc, "7. Лог", 1)
    add_bullets(
        doc,
        [
            "Файл media-monitor_log.txt рядом с exe.",
            "Пишется то же, что видно в консоли.",
            "Если лог больше 10 МБ, старая часть обрезается автоматически.",
        ],
    )

    add_heading(doc, "8. Типичные сообщения", 1)
    add_bullets(
        doc,
        [
            "[ssl] … — у сайта проблема с сертификатом; программа повторит запрос без проверки SSL.",
            "[skip auth] — нужна авторизация, данных нет или они не подошли.",
            "[warn] cannot expand … — не удалось взять список ссылок со страницы.",
            "[error] … — ошибка загрузки конкретной страницы.",
        ],
    )

    add_heading(doc, "9. Советы", 1)
    add_bullets(
        doc,
        [
            "В sites.txt указывайте главную сайта или раздел новостей.",
            "Для точного поиска пишите фразу целиком в words.txt.",
            "После обновления exe ваши txt-конфиги не затираются.",
            "Образец настроек смотрите в settings.example.txt.",
        ],
    )

    add_heading(doc, "10. Сколько времени занимает проверка", 1)
    add_para(
        doc,
        "Программа работает последовательно (сайт за сайтом, страница за страницей). "
        "На каждый сайт обычно приходится: главная + несколько страниц поиска + найденные статьи.",
    )
    add_bullets(
        doc,
        [
            "Ориентир на 1 сайт и 1–2 слова: примерно 1–5 минут (зависит от скорости сайта и числа статей).",
            "Ориентир на 130 сайтов и 1–3 слова: обычно от 2–4 часов до 6–10 часов.",
            "Если много слов и max_scan_urls=0 (без лимита), время может вырасти сильнее.",
            "Медленные сайты, SSL-ошибки и таймауты (до ~25 сек на страницу) удлиняют прогон.",
        ],
    )

    add_para(doc, "Текущий релиз: v1.0.2", bold=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
